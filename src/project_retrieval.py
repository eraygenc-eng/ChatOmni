import json
import math
import re
from functools import lru_cache
from pathlib import Path, PurePosixPath
from docx import Document
from src.project_files import get_project_files


PROJECTS_DATA_DIR = Path(
    "projects_data"
)


PROJECT_TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".jsonc",
    ".xml",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".conf",
    ".log",
    ".properties",
    ".html",
    ".htm",
    ".css",
    ".scss",
    ".sass",
    ".less",
    ".js",
    ".jsx",
    ".mjs",
    ".cjs",
    ".ts",
    ".tsx",
    ".py",
    ".pyw",
    ".java",
    ".kt",
    ".kts",
    ".scala",
    ".c",
    ".h",
    ".cpp",
    ".cc",
    ".cxx",
    ".hpp",
    ".cs",
    ".vb",
    ".fs",
    ".fsx",
    ".go",
    ".rs",
    ".swift",
    ".dart",
    ".php",
    ".rb",
    ".sh",
    ".bash",
    ".zsh",
    ".fish",
    ".ps1",
    ".bat",
    ".cmd",
    ".sql",
    ".graphql",
    ".gql",
    ".proto",
    ".vue",
    ".svelte",
    ".r",
    ".lua",
    ".pl",
    ".pm",
    ".ex",
    ".exs",
    ".erl",
    ".hrl",
    ".clj",
    ".cljs",
    ".cljc",
    ".groovy",
    ".gradle",
    ".tex",
    ".bib",
    ".asm",
    ".s",
    ".ipynb",

    # Useful project-specific text file.
    ".env",

    # Word documents.
    ".docx",
}


PROJECT_TEXT_FILENAMES = {
    "dockerfile",
    "makefile",
    "procfile",
    ".gitignore",
    ".dockerignore",
    ".env.example",
}


CHUNK_MAX_CHARS = 6000
CHUNK_MAX_LINES = 140
CHUNK_OVERLAP_LINES = 12

SEARCH_RESULT_LIMIT = 6
DEEP_BATCH_SIZE = 4

# Maximum size of a single readable source/text file.
MAX_TEXT_FILE_SIZE = 2 * 1024 * 1024  # 2 MB

# DOCX files may contain images and ZIP metadata,
# so allow a larger physical file size.
MAX_DOCX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

# Normal project questions should never scan an
# unlimited number of files.
MAX_GENERIC_SCAN_FILES = 200
MAX_GENERIC_SCAN_BYTES = 25 * 1024 * 1024  # 25 MB

# Safety limit for a whole-project deep review.
MAX_DEEP_SCAN_BYTES = 50 * 1024 * 1024  # 50 MB


IGNORED_PROJECT_DIRS = {
    ".git",
    ".svn",
    ".hg",
    "node_modules",
    ".venv",
    "venv",
    "__pycache__",
    ".pytest_cache",
    ".mypy_cache",
    ".ruff_cache",
    ".cache",
    ".next",
    ".nuxt",
    "dist",
    "build",
    "coverage",
    "target",
    "vendor",
    "site-packages",
    "obj",
}


SENSITIVE_PROJECT_FILENAMES = {
    ".env",
    ".env.local",
    ".env.development",
    ".env.production",
    ".env.test",
    "credentials.json",
    "service-account.json",
    "service_account.json",
}


IMPORTANT_PROJECT_FILENAMES = {
    "readme",
    "readme.md",
    "package.json",
    "pyproject.toml",
    "requirements.txt",
    "dockerfile",
    "docker-compose.yml",
    "compose.yml",
    "main.py",
    "app.py",
    "api.py",
    "index.js",
    "index.ts",
    "vite.config.js",
    "vite.config.ts",
}


STOP_WORDS = {
    "the",
    "a",
    "an",
    "is",
    "are",
    "of",
    "to",
    "for",
    "in",
    "on",
    "with",
    "from",
    "and",
    "or",
    "this",
    "that",
    "file",
    "project",
    "code",

    "bu",
    "şu",
    "bir",
    "ve",
    "veya",
    "ile",
    "için",
    "dosya",
    "dosyada",
    "proje",
    "projede",
    "kod",
    "nasıl",
    "neden",
    "nedir",
}


DEEP_REQUEST_TERMS = {
    "deeply",
    "in depth",
    "in-depth",
    "comprehensive",
    "thorough",
    "entire file",
    "whole file",
    "entire folder",
    "whole folder",

    "derinlemesine",
    "detaylı incele",
    "detaylı analiz",
    "baştan sona",
    "tüm dosyayı",
    "bütün dosyayı",
    "dosyanın tamamını",
    "tüm klasörü",
    "bütün klasörü",
}


def get_search_terms(
    text: str,
):

    words = re.findall(
        r"[\w./\\-]+",
        text.casefold(),
    )

    return {
        word
        for word in words
        if (
            len(word) >= 2
            and
            word not in STOP_WORDS
        )
    }


def is_deep_request(
    question: str,
):

    normalized = (
        question.casefold()
    )

    return any(
        term in normalized
        for term in DEEP_REQUEST_TERMS
    )


def is_text_project_file(file_data: dict):
    extension = (
        file_data.get("extension") or ""
    ).casefold()

    file_name = (
        file_data["file_name"]
    ).casefold()

    relative_path = (
        file_data["relative_path"]
        .replace("\\", "/")
    )

    path_parts = {
        part.casefold()
        for part in PurePosixPath(relative_path).parts[:-1]
    }

    # Ignore generated/dependency/cache directories.
    if path_parts & IGNORED_PROJECT_DIRS:
        return False

    # Never send real environment/credential files
    # into project retrieval.
    if (
        file_name in SENSITIVE_PROJECT_FILENAMES
        or (
            file_name.startswith(".env.")
            and file_name != ".env.example"
        )
    ):
        return False

    return (
        extension in PROJECT_TEXT_EXTENSIONS
        or file_name in PROJECT_TEXT_FILENAMES
    )


def get_file_read_limit(file_data: dict):
    extension = (
        file_data.get("extension") or ""
    ).casefold()

    if extension == ".docx":
        return MAX_DOCX_FILE_SIZE

    return MAX_TEXT_FILE_SIZE


def get_readable_size_bytes(files: list[dict]):
    total = 0

    for file_data in files:
        size_bytes = int(
            file_data.get("size_bytes") or 0
        )

        if size_bytes <= get_file_read_limit(file_data):
            total += size_bytes

    return total


def select_generic_project_files(
    files: list[dict],
    question: str,
):
    search_terms = get_search_terms(question)

    ranked_files = []

    for index, file_data in enumerate(files):
        
        relative_path = (
            file_data["relative_path"]
            .replace("\\", "/")
            .casefold()
        )

        file_name = (
            file_data["file_name"]
            .casefold()
        )

        score = 0

        # Prefer important entry/configuration files.
        if file_name in IMPORTANT_PROJECT_FILENAMES:
            score += 4

        # Prefer filenames/paths related to the question.
        for term in search_terms:
            if term in file_name:
                score += 12
            elif term in relative_path:
                score += 6

        ranked_files.append(
            (
                score,
                -index,
                file_data,
            )
        )

    ranked_files.sort(
        key=lambda item: (
            item[0],
            item[1],
        ),
        reverse=True,
    )

    selected_files = []
    total_bytes = 0

    for _, _, file_data in ranked_files:
        size_bytes = int(
            file_data.get("size_bytes") or 0
        )

        if size_bytes > get_file_read_limit(file_data):
            continue

        if len(selected_files) >= MAX_GENERIC_SCAN_FILES:
            break

        if (
            total_bytes + size_bytes
            > MAX_GENERIC_SCAN_BYTES
        ):
            continue

        selected_files.append(file_data)
        total_bytes += size_bytes

    return selected_files


def get_safe_project_file_path(
    project_id: str,
    relative_path: str,
):

    project_root = (
        PROJECTS_DATA_DIR /
        project_id
    ).resolve()

    normalized_path = PurePosixPath(
        relative_path.replace(
            "\\",
            "/",
        )
    )

    file_path = (
        project_root /
        Path(
            *normalized_path.parts
        )
    ).resolve()

    try:

        file_path.relative_to(
            project_root
        )

    except ValueError:

        return None

    return file_path


def read_project_file(
    project_id: str,
    relative_path: str,
):

    file_path = (
        get_safe_project_file_path(
            project_id,
            relative_path,
        )
    )

    if (
        file_path is None
        or
        not file_path.exists()
        or
        not file_path.is_file()
    ):
        return None

    suffix = (
        file_path
        .suffix
        .lower()
    )

    max_file_size = (
        MAX_DOCX_FILE_SIZE
        if suffix == ".docx"
        else MAX_TEXT_FILE_SIZE
    )

    try:
        if file_path.stat().st_size > max_file_size:
            return None

    except OSError:
        return None

    # DOCX files are binary ZIP-based documents,
    # so they must be parsed with python-docx.
    if suffix == ".docx":

        try:

            document = Document(
                file_path
            )

            text_parts = []

            # Extract normal paragraphs.
            for paragraph in document.paragraphs:

                text = (
                    paragraph
                    .text
                    .strip()
                )

                if text:
                    text_parts.append(
                        text
                    )

            # Extract table contents as well.
            for table in document.tables:

                for row in table.rows:

                    row_values = [
                        cell.text.strip()
                        for cell in row.cells
                    ]

                    if any(row_values):

                        text_parts.append(
                            " | ".join(
                                row_values
                            )
                        )

            document_text = (
                "\n".join(
                    text_parts
                )
                .strip()
            )

            if not document_text:
                return None

            return document_text

        except Exception:

            return None

    # Normal source-code and text files.
    try:

        return file_path.read_text(
            encoding="utf-8"
        )

    except UnicodeDecodeError:

        try:

            return file_path.read_text(
                encoding="utf-8-sig"
            )

        except UnicodeDecodeError:

            return None


def chunk_file_content(
    relative_path: str,
    content: str,
):

    lines = content.splitlines()

    if not lines:

        return []

    chunks = []

    start = 0

    while start < len(lines):

        end = start
        character_count = 0

        while (
            end < len(lines)
            and
            end - start <
            CHUNK_MAX_LINES
        ):

            next_length = (
                len(lines[end])
                + 1
            )

            if (
                character_count > 0
                and
                character_count
                + next_length
                >
                CHUNK_MAX_CHARS
            ):

                break

            character_count += (
                next_length
            )

            end += 1

        if end == start:

            end = start + 1

        chunk_text = "\n".join(
            lines[
                start:end
            ]
        )

        chunks.append(
            {
                "path":
                    relative_path,

                "start_line":
                    start + 1,

                "end_line":
                    end,

                "content":
                    chunk_text,
            }
        )

        if end >= len(lines):

            break

        start = max(
            end
            - CHUNK_OVERLAP_LINES,
            start + 1,
        )

    return chunks


def build_chunks(
    project_id: str,
    files: list[dict],
):

    chunks = []
    skipped_files = []

    for file_data in files:

        relative_path = (
            file_data[
                "relative_path"
            ]
        )

        content = read_project_file(
            project_id,
            relative_path,
        )

        if content is None:

            skipped_files.append(
                relative_path
            )

            continue

        chunks.extend(
            chunk_file_content(
                relative_path,
                content,
            )
        )

    return (
        chunks,
        skipped_files,
    )

def get_file_signature(files: list[dict]):
    return tuple(
        (
            file_data["relative_path"],
            int(
                file_data.get("size_bytes") or 0
            ),
            str(
                file_data.get("updated_at") or ""
            ),
        )
        for file_data in files
    )


@lru_cache(maxsize=2)
def build_chunks_cached(
    project_id: str,
    file_signature: tuple,
):
    files = [
        {
            "relative_path": relative_path,
        }
        for (
            relative_path,
            _,
            _,
        ) in file_signature
    ]

    return build_chunks(
        project_id,
        files,
    )


def get_cached_chunks(
    project_id: str,
    files: list[dict],
):
    return build_chunks_cached(
        project_id,
        get_file_signature(files),
    )


def detect_file_from_question(
    project_files: list[dict],
    question: str,
):

    normalized_question = (
        question.casefold()
    )

    matches = []

    for file_data in project_files:

        relative_path = (
            file_data[
                "relative_path"
            ]
            .casefold()
        )

        file_name = (
            file_data[
                "file_name"
            ]
            .casefold()
        )

        if (
            relative_path
            in normalized_question
            or
            file_name
            in normalized_question
        ):

            matches.append(
                file_data
            )

    if len(matches) == 1:

        return matches[0]

    return None


def resolve_target_file(
    project_files: list[dict],
    target: str,
    question: str,
):

    clean_target = (
        target.strip()
        .strip('"')
        .strip("'")
        .replace(
            "\\",
            "/",
        )
        .casefold()
    )

    if not clean_target:

        detected = (
            detect_file_from_question(
                project_files,
                question,
            )
        )

        return detected

    exact_path_matches = [
        file_data
        for file_data in project_files
        if (
            file_data[
                "relative_path"
            ]
            .replace(
                "\\",
                "/",
            )
            .casefold()
            ==
            clean_target
        )
    ]

    if len(
        exact_path_matches
    ) == 1:

        return (
            exact_path_matches[0]
        )

    file_name_matches = [
        file_data
        for file_data in project_files
        if (
            file_data[
                "file_name"
            ]
            .casefold()
            ==
            clean_target
        )
    ]

    if len(
        file_name_matches
    ) == 1:

        return (
            file_name_matches[0]
        )

    suffix_matches = [
        file_data
        for file_data in project_files
        if (
            file_data[
                "relative_path"
            ]
            .replace(
                "\\",
                "/",
            )
            .casefold()
            .endswith(
                f"/{clean_target}"
            )
        )
    ]

    if len(
        suffix_matches
    ) == 1:

        return (
            suffix_matches[0]
        )

    return None


def resolve_scope_files(
    project_files: list[dict],
    target: str,
):

    clean_target = (
        target.strip()
        .strip("/")
        .strip("\\")
        .replace(
            "\\",
            "/",
        )
        .casefold()
    )

    if not clean_target:

        return []

    matches = []

    for file_data in project_files:

        relative_path = (
            file_data[
                "relative_path"
            ]
            .replace(
                "\\",
                "/",
            )
            .casefold()
        )

        path_parts = (
            PurePosixPath(
                relative_path
            )
            .parts
        )

        if (
            relative_path.startswith(
                f"{clean_target}/"
            )
            or
            clean_target
            in path_parts[:-1]
        ):

            matches.append(
                file_data
            )

    return matches


def score_chunk(
    chunk: dict,
    search_terms: set[str],
):

    if not search_terms:

        return 0

    path = (
        chunk["path"]
        .casefold()
    )

    content = (
        chunk["content"]
        .casefold()
    )

    score = 0

    for term in search_terms:

        if term in path:

            score += 8

        occurrence_count = (
            content.count(
                term
            )
        )

        score += min(
            occurrence_count,
            5,
        )

    return score


def select_relevant_chunks(
    chunks: list[dict],
    question: str,
):

    search_terms = (
        get_search_terms(
            question
        )
    )

    scored_chunks = [
        (
            score_chunk(
                chunk,
                search_terms,
            ),
            chunk,
        )

        for chunk in chunks
    ]

    scored_chunks.sort(
        key=lambda item: item[0],
        reverse=True,
    )

    positive_chunks = [
        chunk
        for score, chunk
        in scored_chunks
        if score > 0
    ]

    if positive_chunks:

        return positive_chunks[
            :SEARCH_RESULT_LIMIT
        ]

    return chunks[
        :SEARCH_RESULT_LIMIT
    ]


def build_result(
    mode: str,
    items: list[dict],
    target: str | None = None,
    skipped_files: list[str] | None = None,
    batch: int = 0,
    total_batches: int = 1,
):

    has_more = (
        batch
        <
        total_batches - 1
    )

    return json.dumps(
        {
            "status":
                "ok",

            "mode":
                mode,

            "target":
                target,

            "batch":
                batch,

            "total_batches":
                total_batches,

            "has_more":
                has_more,

            "next_batch":
                (
                    batch + 1
                    if has_more
                    else None
                ),

            "items":
                items,

            "skipped_files":
                skipped_files
                or [],
        },
        ensure_ascii=False,
    )


def build_batched_result(
    mode: str,
    chunks: list[dict],
    batch: int,
    target: str | None = None,
    skipped_files: list[str] | None = None,
):

    if not chunks:

        return json.dumps(
            {
                "status":
                    "error",

                "message":
                    "No readable project content was found.",
            },
            ensure_ascii=False,
        )

    total_batches = (
        math.ceil(
            len(chunks)
            /
            DEEP_BATCH_SIZE
        )
    )

    if (
        batch < 0
        or
        batch >= total_batches
    ):

        return json.dumps(
            {
                "status":
                    "error",

                "message":
                    "Invalid project review batch.",
            },
            ensure_ascii=False,
        )

    start = (
        batch
        *
        DEEP_BATCH_SIZE
    )

    end = (
        start
        +
        DEEP_BATCH_SIZE
    )

    return build_result(
        mode=mode,
        target=target,
        items=chunks[
            start:end
        ],
        skipped_files=
            skipped_files,
        batch=batch,
        total_batches=
            total_batches,
    )


def retrieve_project_context(
    project_id: str,
    question: str,
    mode: str,
    target: str = "",
    batch: int = 0,
):

    project_files = (
        get_project_files(
            project_id
        )
    )

    text_files = [
        file_data
        for file_data
        in project_files
        if is_text_project_file(
            file_data
        )
    ]

    if not text_files:

        return json.dumps(
            {
                "status":
                    "error",

                "message":
                    "The active project has no readable text or source files.",
            },
            ensure_ascii=False,
        )

    # Targeted
    if mode == "targeted":

        target_file = (
            resolve_target_file(
                text_files,
                target,
                question,
            )
        )

        if target_file is not None:

            selected_files = [
                target_file
            ]

            chunks, skipped = (
                get_cached_chunks(
                    project_id,
                    selected_files,
                )
            )

            if is_deep_request(
                question
            ):

                return (
                    build_batched_result(
                        mode="targeted",
                        chunks=chunks,
                        batch=batch,
                        target=
                            target_file[
                                "relative_path"
                            ],
                        skipped_files=
                            skipped,
                    )
                )

            relevant_chunks = (
                select_relevant_chunks(
                    chunks,
                    question,
                )
            )

            return build_result(
                mode="targeted",
                target=
                    target_file[
                        "relative_path"
                    ],
                items=
                    relevant_chunks,
                skipped_files=
                    skipped,
            )

        # No explicit file:
        # focused search across a bounded set of project files.
        selected_files = (
            select_generic_project_files(
                text_files,
                question,
            )
        )

        if not selected_files:
            return json.dumps(
                {
                    "status": "error",
                    "message":
                        "No safe readable project files were found for this search.",
                },
                ensure_ascii=False,
            )

        chunks, skipped = (
            get_cached_chunks(
                project_id,
                selected_files,
            )
        )

        relevant_chunks = (
            select_relevant_chunks(
                chunks,
                question,
            )
        )

        return build_result(
            mode="targeted",
            items=
                relevant_chunks,
            skipped_files=
                skipped,
        )
    
    # Scoped
    if mode == "scoped":

        scoped_files = (
            resolve_scope_files(
                text_files,
                target,
            )
        )

        if not scoped_files:

            return json.dumps(
                {
                    "status":
                        "error",

                    "message":
                        (
                            "No files were found "
                            "for the requested project scope."
                        ),
                },
                ensure_ascii=False,
            )

        chunks, skipped = (
            get_cached_chunks(
                project_id,
                scoped_files,
            )
        )

        if is_deep_request(
            question
        ):

            return (
                build_batched_result(
                    mode="scoped",
                    chunks=chunks,
                    batch=batch,
                    target=target,
                    skipped_files=
                        skipped,
                )
            )

        relevant_chunks = (
            select_relevant_chunks(
                chunks,
                question,
            )
        )

        return build_result(
            mode="scoped",
            target=target,
            items=
                relevant_chunks,
            skipped_files=
                skipped,
        )

    # Deep Project Review
    if mode == "deep":

        readable_size = (
            get_readable_size_bytes(
                text_files
            )
        )

        if readable_size > MAX_DEEP_SCAN_BYTES:
            return json.dumps(
                {
                    "status": "error",
                    "message": (
                        "This project is too large for a safe whole-project "
                        "deep review. Please review a specific folder or "
                        "module instead."
                    ),
                },
                ensure_ascii=False,
            )

        chunks, skipped = (
            get_cached_chunks(
                project_id,
                text_files,
            )
        )

        return (
            build_batched_result(
                mode="deep",
                chunks=chunks,
                batch=batch,
                target="whole-project",
                skipped_files=
                    skipped,
            )
        )

    return json.dumps(
        {
            "status":
                "error",

            "message":
                "Unknown project search mode.",
        },
        ensure_ascii=False,
    )